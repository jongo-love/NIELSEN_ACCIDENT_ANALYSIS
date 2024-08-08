# LINE 1-10 I'M CREATING ROUTES FOR THE HOME PAGE.
import os
from geopy.exc import GeocoderTimedOut
import time
from flask_mail import Message,Mail
from flask import Flask, flash, redirect, render_template, request, url_for, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_wtf import FlaskForm
from wtforms import EmailField, StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired,EqualTo
from werkzeug.security import generate_password_hash, check_password_hash
from itsdangerous import URLSafeTimedSerializer
from flask_login import LoginManager, UserMixin, login_user, logout_user, current_user, login_required

#IMPORTING LIBRARIES THAT WILL BE USED IN DATA SCIENCE.
import sqlite3
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import seaborn as sns    
import seaborn as sns    
import matplotlib
from matplotlib import cm
from matplotlib import cm
import matplotlib.patches as mpatches
import geopandas as gpd
from geopy.geocoders import Nominatim
from shapely.geometry import Point
from docx import Document
from docx.shared import Inches
matplotlib.use('Agg')

app = Flask(__name__)
login_manager = LoginManager(app)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///nielsenaccident.db'  # SQLite database file path
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'jongolove01@gmail.com'
app.config['MAIL_PASSWORD'] = 'ctsy mhos lokr jzdn'
app.config['MAIL_DEFAULT_SENDER'] = 'jongolove01@gmail.com'
db = SQLAlchemy(app)
mail = Mail(app)

# Define the User model/ Creating the USER table in the database.
class User(UserMixin,db.Model):
    id = db.Column(db.Integer, primary_key=True,)
    username = db.Column(db.String(20), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)

class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    email = EmailField('Email', validators=[DataRequired()])
    submit = SubmitField('Login')


class RegistrationForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    email = EmailField('Email', validators=[DataRequired()])
    submit = SubmitField('Register')

class ForgotPasswordForm(FlaskForm):
    email = EmailField('Email', validators=[DataRequired()])
    submit = SubmitField('Submit')

class ResetPasswordForm(FlaskForm):
    password = PasswordField('New Password', validators=[DataRequired()])
    confirm_password = PasswordField('Confirm New Password', validators=[DataRequired(), EqualTo('password')])
    submit = SubmitField('Reset Password')

# Configure Flask-Login
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

#with app.app_context():
    #db.create_all()
    #db.session.commit()

@app.route('/home')
@login_required
def HOME():
    return render_template('home.html')

@app.route('/about')
@login_required
def ABOUT():
    return render_template('about.html')

@app.route('/contact',methods=['GET', 'POST'])
@login_required
def CONTACT():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        message = request.form['message']

        # Compose the email message
        sender = app.config['MAIL_DEFAULT_SENDER']  # Assigning the sender
        msg = Message(subject='New Contact Form Submission',
                    sender=sender,
                    recipients=['jongolove01@gmail.com'])
        msg.body = f'Name: {name}\nEmail: {email}\nMessage: {message}'

        # Send the email
        mail.send(msg)

        # Redirect to a thank you page or home page after successful submission
        return redirect(url_for('THANKYOU'))
    return render_template('contact.html')

@app.route('/thank_you')
@login_required
def THANKYOU():
    return render_template('thanks.html')

@app.route('/', methods=['GET', 'POST'])
def LOGIN():
    time.sleep(1.5)
    form = LoginForm()
    error = None
    if form.validate_on_submit():
        # You can handle the form data here
        username = form.username.data
        password = form.password.data
        email = form.email.data
        user = User.query.filter_by(username=username).first()

        if user:
            if check_password_hash(user.password, password)and user.email == email:
                login_user(user)
                flash('Logged in successfully.')
                return redirect(url_for('HOME'))
            else:
                error = 'Incorrect password. Please try again.'
        else:
            error = 'Username not found. Please register.'
    return render_template('login.html', form=form, error=error)

@app.route('/logout')
@login_required
def LOGOUT():
    logout_user()
    flash('Logged out successfully.')
    return redirect(url_for('LOGIN'))

@app.route('/register', methods=['GET', 'POST']) 
def REGISTER():
    time.sleep(1.5)
    form = RegistrationForm()
    error = None
    if form.validate_on_submit():
        # You can handle the form data here
        username = form.username.data
        password = form.password.data
        email = form.email.data
        existing_user = User.query.filter_by(username=username).first()

        if existing_user:
            error = 'Username already exists. Please choose a different one.'
        else:
            hashed_password = generate_password_hash(password)
            new_user = User(username=username, password=hashed_password,email=email)
            db.session.add(new_user)
            db.session.commit()
            return redirect(url_for('LOGIN'))
    
    return render_template('register.html', form=form, error=error)

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    form = ForgotPasswordForm()
    if form.validate_on_submit():
        email = form.email.data
        user = User.query.filter_by(email=email).first()
        if user:
            # Generate a password reset token
            serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
            token = serializer.dumps(user.email, salt='password-reset-salt')
            # Send password reset email with the token
            reset_url = url_for('reset_password', token=token, _external=True)
            msg = Message('Password Reset', recipients=[user.email])
            msg.body = f'Click the following link to reset your password: {reset_url}'
            mail.send(msg)  # Send the email
            flash('Password reset email sent. Please check your email.')
            return redirect(url_for('LOGIN'))
        else:
            flash('Email address not found. Please try again.')
    return render_template('forgot_password.html', form=form)


@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    serializer = URLSafeTimedSerializer(app.config['SECRET_KEY'])
    try:
        email = serializer.loads(token, salt='password-reset-salt', max_age=3600)  # Token expires after 1 hour
        user = User.query.filter_by(email=email).first()
        if not user:
            flash('Invalid or expired token. Please try again.')
            return redirect(url_for('forgot_password'))
        form = ResetPasswordForm()
        if form.validate_on_submit():
            # Update user's password
            hashed_password = generate_password_hash(form.password.data)
            user.password = hashed_password
            db.session.commit()
            flash('Your password has been reset. You can now log in with your new password.')
            return redirect(url_for('LOGIN'))
        return render_template('reset_password.html', form=form)
    except Exception as e:
        print(e)
        flash('Invalid or expired token. Please try again.')
        return redirect(url_for('forgot_password'))

#the function of this route is, when the Data Analysis and Visualization service is clicked, the page should return a beautiful
# dashboard which WILL have buttons for the different kind of visualizations.
#@app.route('/data_analysis')
#def DASHBOARD():
    #return render_template('dashboard.html')


@app.route('/data_analysis')
@login_required
def DATALYSIS():
    conn = sqlite3.connect('C:\\Users\\user\\NIELSEN_ACCIDENT_ANALYSIS\\instance\\nielsenaccident.db')
    conn = sqlite3.connect('C:\\Users\\user\\NIELSEN_ACCIDENT_ANALYSIS\\instance\\nielsenaccident.db')
    df = pd.read_sql_query("SELECT * FROM accidents", conn)
    conn.close()
    
    df['Start_Time'] = pd.to_datetime(df['Start_Time'])
    df['End_Time'] = pd.to_datetime(df['End_Time'])
    
    #CITY ANALYSIS.
    #CITY ANALYSIS.
    #CITY ANALYSIS.

    city_accidents = df['City'].value_counts().reset_index()
    city_accidents.columns = ['City', 'Accident Cases']
    
    top_8_cities = city_accidents.head(8)
    
    # Customized visualization
    ax = plt.subplots(figsize=(12, 7), dpi=80)
    ax = plt.subplots(figsize=(12, 7), dpi=80)
    cmap = cm.get_cmap('rainbow', 10)
    clrs = [cm.colors.rgb2hex(cmap(i)) for i in range(cmap.N)]
    
    ax = sns.barplot(y=top_8_cities['Accident Cases'], x=top_8_cities['City'], palette=clrs)
    
    total = sum(city_accidents['Accident Cases'])
    for i in ax.patches:
        ax.text(i.get_x() + 0.03, i.get_height() - 2500,
                f"{round((i.get_height()/total)*100, 2)}%", fontsize=15, weight='bold', color='white')
    
    plt.title('Top 10 Cities in US with the Most Number of Road Accident Cases (2016-2020)',
              size=20, color='grey')
    plt.ylim(0, 400)
    plt.xticks(rotation=10, fontsize=12)
    plt.yticks(fontsize=12)
    
    ax.set_xlabel('\nCities\n', fontsize=15, color='grey')
    ax.set_ylabel('Accident Cases\n', fontsize=15, color='grey')
    
    for spine in ['bottom', 'left']:
        ax.spines[spine].set_color('white')
        ax.spines[spine].set_linewidth(1.5)
    
    right_side = ax.spines['right']
    right_side.set_visible(False)
    top_side = ax.spines['top']
    top_side.set_visible(False)
    
    ax.set_axisbelow(True)
    ax.grid(color='#b2d6c7', linewidth=1, axis='y', alpha=0.3)
    
    MA = mpatches.Patch(color=clrs[0], label='City with Maximum\n no. of Road Accidents')
    ax.legend(handles=[MA], prop={'size': 10.5}, loc='best', borderpad=1, 
              labelcolor=clrs[0], edgecolor='white')

    custom_plot_path = 'static/custom_plot.png'
    plt.savefig(custom_plot_path)
    plt.close()

    #STATE ANALYSIS.
    #STATE ANALYSIS.
    #STATE ANALYSIS.

    #Dictionary of US state codes and their corresponding names
    us_states = {'MA': 'Massachusetts',
        'CA': 'California',
        'UT': 'Utah',
        'NC': 'North Carolina',
        'TX': 'Texas',
        'IL': 'Illinois',
        'IN': 'Indiana',
        'VA': 'Virginia',
        'AZ': 'Arizona',
        'NY': 'New York',
        'LA': 'Louisiana',
        'AL': 'Alabama',
        'SC': 'South Carolina',
        'MI': 'Michigan',
        'OK': 'Oklahoma',
        'FL': 'Florida',
        'OH': 'Ohio',
        'TN': 'Tennessee',
        'MN': 'Minnesota',
        'PA': 'Pennsylvania',
        'NH': 'New Hampshire',
        'RI': 'Rhode Island',
        'NM': 'New Mexico',
        'KY': 'Kentucky',
        'MO': 'Missouri',
        'WA': 'Washington',
        'ME': 'Maine',
        'GA': 'Georgia',
        'NJ': 'New Jersey',
        'OR': 'Oregon',
        'MD': 'Maryland',
        'NE': 'Nebraska',
        'CT': 'Connecticut',
        'IA': 'Iowa',
        'CO': 'Colorado',
        'WI': 'Wisconsin',
        'DE': 'Delaware',
        'NV': 'Nevada',
        'MS': 'Mississippi',
        'WV': 'West Virginia',
        'DC': 'District of Columbia',
        'AR': 'Arkansas',
        'KS': 'Kansas'}

    # Count the number of accident cases per state
    state = df['State'].value_counts().reset_index()
    state.columns =['State', 'Cases']

    # Convert state codes to state names
    state['State'] = state['State'].apply(lambda x: us_states.get(x, x))

    # Get the top 10 states by accident cases
    top_10_states = state.head(10)

    # Create the barplot and line plot
    fig, ax = plt.subplots(figsize=(12, 8), dpi=80)
    cmap = cm.get_cmap('winter', 10)
    clrs = [cm.colors.rgb2hex(cmap(i)) for i in range(cmap.N)]

    ax = sns.barplot(y=top_10_states['Cases'], x=top_10_states['State'], palette=clrs)
    ax1 = ax.twinx()
    sns.lineplot(data=top_10_states, marker='o', x='State', y='Cases', color='white', alpha=.8, ax=ax1)

    # Annotate the barplot with case numbers and percentages
    total = df.shape[0]
    for i in ax.patches:
        ax.text(i.get_x() + 0.03, i.get_height() - 2500,
                f"{int(i.get_height()):,}\n({round((i.get_height() / total) * 100, 1)}%)",
                fontsize=15, weight='bold', color='white')
        
        # Customize the plot's appearance
    plt.title('Top 10 US States with the Most Number of Road Accident Cases (2016-2020)', size=20, color='grey')
    ax.set_xlabel('\nStates\n', fontsize=15, color='grey')
    ax.set_ylabel('\nAccident Cases\n', fontsize=15, color='grey')
    ax1.axes.yaxis.set_visible(False)
    plt.xticks(rotation=10, fontsize=12)
    plt.yticks(fontsize=12)
    plt.ylim(0, 500000)
    ax.grid(color='#b2d6c7', linewidth=1, axis='y', alpha=0.3)

    # Customize spines
    for spine in ['bottom', 'left']:
        ax.spines[spine].set_color('white')
        ax.spines[spine].set_linewidth(1.5)
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)

    # Create and add legend
    MA = mpatches.Patch(color=clrs[0], label='State with Maximum\n no. of Road Accidents')
    ax.legend(handles=[MA], prop={'size': 10.5}, loc='best', borderpad=1, labelcolor=clrs[0], edgecolor='white')

    state_plot_path = 'static/state_plot.png'
    plt.savefig(state_plot_path)
    plt.close()

    #Map analysis
    #Map analysis
    #Map analysis
    # Initialize the geolocator
    # Initialize the geolocator
    geolocator = Nominatim(user_agent="Your_Name")

    def get_lat_lng(city, retries=3, wait_time=2):
        for i in range(retries):
            try:
                location = geolocator.geocode(city, timeout=100)
                return (location.latitude, location.longitude) if location else (None, None)
            except GeocoderTimedOut:
                if i < retries - 1:
                    time.sleep(wait_time)
                    continue
                else:
                    return (None, None)

    # Get the latitude and longitude for the top 10 cities
    top_10_cities = city_accidents.head(10)['City']
    city_coordinates = {city: get_lat_lng(city) for city in top_10_cities}

    # Filter out cities where coordinates could not be found
    city_coordinates = {city: coords for city, coords in city_coordinates.items() if None not in coords}

    # Create a DataFrame for the top 10 cities with their coordinates
    top_10_cities_df = df[df['City'].isin(city_coordinates.keys())]
    top_10_cities_df['Latitude'] = top_10_cities_df['City'].map(lambda city: city_coordinates[city][0])
    top_10_cities_df['Longitude'] = top_10_cities_df['City'].map(lambda city: city_coordinates[city][1])

    # Create a GeoDataFrame
    geometry = [Point(xy) for xy in zip(top_10_cities_df['Longitude'], top_10_cities_df['Latitude'])]
    geo_df = gpd.GeoDataFrame(top_10_cities_df, geometry=geometry)

    # Load the US state boundaries
    world = gpd.read_file('ne_110m_admin_0_countries.shp')
    us_states = world[world['NAME'] == 'United States of America']

    # Plotting
    fig, ax = plt.subplots(figsize=(15, 15))
    ax.set_xlim([-125, -65])
    ax.set_ylim([22, 55])
    us_states.boundary.plot (ax=ax, color='grey', linewidth=1)

    # Define colors and sizes for the top 10 cities
    colors = ['#e6194B', '#f58231', '#ffe119', '#bfef45', '#3cb44b', '#aaffc3', '#42d4f4', '#4363d8', '#911eb4', '#f032e6']
    sizes = [100 + (i * 50) for i in range(10)]

    # Plot each city with a different color and size
    for i, city in enumerate(top_10_cities):
        city_df = geo_df[geo_df['City'] == city]
        city_df.plot(ax=ax, marker='o', color=colors[i], markersize=sizes[i], label=city, alpha=0.7)

    plt.legend(prop={'size': 13}, loc='best', bbox_to_anchor=(1.05, 1), edgecolor='white', title="Cities", title_fontsize=15)
    plt.title('Visualization of Top 10 Accident Prone Cities in US (2016-2020)', size=20, color='grey')
    plt.axis('off')
    plt.show()

    # Save the plot
    map_plot_path = 'static/map_plot.png'
    plt.savefig(map_plot_path, bbox_inches='tight')
    plt.close()


    #Timezone analysis
    #Timezone analysis
    #Timezone analysis

    timezone_df = pd.DataFrame(df['Timezone'].value_counts()).reset_index()
    timezone_df.columns = ['Timezone', 'Cases']

    fig, ax = plt.subplots(figsize=(12, 8), dpi=80)

    cmap = cm.get_cmap('spring', 4)
    clrs = [matplotlib.colors.rgb2hex(cmap(i)) for i in range(cmap.N)]

    ax = sns.barplot(y=timezone_df['Cases'], x=timezone_df['Timezone'], palette='spring')

    total = df.shape[0]
    for i in ax.patches:
        percentage = round(i.get_height() * 100 / total)
        ax.text(i.get_x() + i.get_width() / 2, i.get_height() + total * 0.01,
                '{}%'.format(percentage), fontsize=15, weight='bold',
                color='black', ha='center')

    plt.ylim(0, timezone_df['Cases'].max() + total * 0.05)
    plt.title('Percentage of Accident Cases for Different Timezones in US (2016-2020)', size=20, color='grey')
    plt.ylabel('Accident Cases', fontsize=15, color='grey')
    plt.xlabel('Timezones', fontsize=15, color='grey')
    plt.xticks(fontsize=13)
    plt.yticks(fontsize=12)

    for i in ['top', 'right']:
        side = ax.spines[i]
        side.set_visible(False)

    ax.set_axisbelow(True)
    ax.grid(color='#b2d6c7', linewidth=1, axis='y', alpha=.3)
    ax.spines['bottom'].set_bounds(-0.5, len(timezone_df['Timezone']) - 0.5)
    ax.spines['left'].set_bounds(0, max(timezone_df['Cases']) * 1.05)

    MA = mpatches.Patch(color=clrs[0], label='Timezone with Maximum\n no. of Road Accidents')
    MI = mpatches.Patch(color=clrs[-1], label='Timezone with Minimum\n no. of Road Accidents')
    ax.legend(handles=[MA, MI], prop={'size': 10.5}, loc='best', borderpad=1,
            labelcolor=[clrs[0], 'grey'], edgecolor='white')

    plt.tight_layout()

    timezone_plot_path = 'static/timezone_plot.png'
    plt.savefig(timezone_plot_path)
    plt.close()


    #STREET ANALYSIS
    #STREET ANALYSIS
    #STREET ANALYSIS

    street_df = pd.DataFrame(df['Street'].value_counts()).reset_index()
    street_df.columns = ['Street', 'Cases']

    top_ten_streets_df = street_df.head(10)

    fig, ax = plt.subplots(figsize=(12, 6), dpi=80)

    cmap = cm.get_cmap('gnuplot2', 10)
    clrs = [matplotlib.colors.rgb2hex(cmap(i)) for i in range(cmap.N)]

    ax = sns.barplot(y=top_ten_streets_df['Cases'], x=top_ten_streets_df['Street'], palette='gnuplot2')
    ax1 = ax.twinx()
    sns.lineplot(data=top_ten_streets_df, marker='o', x='Street', y='Cases', color='white', alpha=.8)

    total = df.shape[0]
    for i in ax.patches:
        ax.text(i.get_x() + i.get_width() / 2, i.get_height() + total * 0.01,
                '{:,d}'.format(int(i.get_height())), fontsize=12.5, weight='bold',
                color='black', ha='center')

    ax.set_ylim(0, top_ten_streets_df['Cases'].max() + total * 0.05)
    ax1.set_ylim(0, top_ten_streets_df['Cases'].max() + total * 0.1)
    plt.title('\nTop 10 Accident Prone Streets in US (2016-2020)\n', size=20, color='grey')

    ax1.axes.yaxis.set_visible(False)
    ax.set_xlabel('\nStreet\n', fontsize=15, color='grey')
    ax.set_ylabel('\nAccident Cases\n', fontsize=15, color='grey')

    for i in ['top', 'right']:
        side1 = ax.spines[i]
        side1.set_visible(False)
        side2 = ax1.spines[i]
        side2.set_visible(False)

    ax.set_axisbelow(True)
    ax.grid(color='#b2d6c7', linewidth=1, axis='y', alpha=.3)

    ax.spines['bottom'].set_bounds(-0.5, len(top_ten_streets_df['Street']) - 0.5)
    ax.spines['left'].set_bounds(0, max(top_ten_streets_df['Cases']) * 1.05)
    ax1.spines['bottom'].set_bounds(-0.5, len(top_ten_streets_df['Street']) - 0.5)
    ax1.spines['left'].set_bounds(0, max(top_ten_streets_df['Cases']) * 1.05)
    ax.tick_params(axis='both', which='major', labelsize=12)

    MA = mpatches.Patch(color=clrs[1], label='Street with Maximum\n no. of Road Accidents')
    MI = mpatches.Patch(color=clrs[-2], label='Street with Minimum\n no. of Road Accidents')
    ax.legend(handles=[MA, MI], prop={'size': 10.5}, loc='best', borderpad=1,
            labelcolor=[clrs[1], 'grey'], edgecolor='white')

    plt.tight_layout()
    street_plot_path = 'static/street_plot.png'
    plt.savefig(street_plot_path)
    plt.close()


    # Pass the data and paths to the template
    return render_template('datalysis.html',street_path=street_plot_path ,plot_path=custom_plot_path, state_path=state_plot_path, time_path=timezone_plot_path , map_path=map_plot_path)



#INCLUDE THE DATABASE IMPLEMENTATION FOR THE CODE PRESUME FUNCTIONALITY.

#INSIGHT.
def city_cases_percentage(city_accidents, val, operator):
    if operator == '<':
        res = city_accidents[city_accidents['Cases'] < val].shape[0]
    elif operator == '>':
        res = city_accidents[city_accidents['Cases'] > val].shape[0]
    elif operator == '=':
        res = city_accidents[city_accidents['Cases'] == val].shape[0]
    percentage = round(res * 100 / city_accidents.shape[0], 2)
    print(f'{res} Cities, {percentage}%')
    return f'{res} Cities, {percentage}%'

def create_report(title, insights, image_path, output_path):
    doc = Document()
    # Add title
    doc.add_heading(title, 0)
    
    # Add insights
    doc.add_heading('Insights', level=1)
    for insight in insights:
        doc.add_paragraph(insight)
    
    # Add visualization
    doc.add_heading('Visualization', level=1)
    doc.add_paragraph('Visualization of the data:')
    doc.add_picture(image_path, width=Inches(6))
    
    
    doc.save(output_path)
    print(f"Report saved as {output_path}")

@app.route('/download-report', methods=['GET'])
@login_required
def generate_report():
    # Define visualizations to generate
    visualizations = [
        {
            'title': 'Top 10 Accident-Prone Cities in US',
            'insights': [
                "In this Dataset, we have the records of total 10,657 Cities.",
                "1. 11% (1,167 Cities) cities in US, have only 1 accident record in past 5 years.",
                "2. Around 81% (8,682 Cities) of all cities in US, have less than 100 total no. of road accidents.",
                "3. 97.64% (10,406 Cities) cities in US, have the road accident records (2016-2020), less than 1,000.",
                "4. There are 251 Cities (2.36%) in US, which have more than 1,000 total no. of road accidents in past 5 years.",
                "5. 40 Cities (0.38%) in US, have more than 5,000 road accident records.",
                "6. Only 13 Cities (0.12%) in US, have more than 10,000 road accident records."
            ],
            'image_path': 'static/map_plot.png',  # Ensure this path is correct
            'output_path': 'static/Top_10_Accident_Prone_Cities_Report.docx'
        },
        {
            'title': 'Top 10 Cities in US',
            'insights': [
                "In this Dataset, we have the records of total 10,657 Cities.",
                "Los Angeles is the city with highest (2.64%) no. of road accidents in US (2016-2020)."
                "Miami is the city with 2nd highest (2.39%) no. of road accidents in US (2016-2020)."
                "Around 14%' accident records of past 5 years are only from these 10 cities out of 10,657 cities in US (as per the dataset)."
                "In past 5 years (2016-2020) yearly 7,997 road accidents (average) happened in Los Angeles."
                "In Los Angeles averagely in every 12 hours 11 accidents occurred"
            ],
            'image_path': 'static/custom_plot.png',  # Ensure this path is correct
            'output_path': 'static/Top_10_Accident Cities_Report.docx'
        },
        #add more reports.
    ]

    # Generate reports
    for viz in visualizations:
        create_report(viz['title'], viz['insights'], viz['image_path'], viz['output_path'])


@app.route('/download-map-report', methods=['GET'])
@login_required
def download_map_report():
    map_report_path = 'static/Top_10_Accident_Prone_Cities_Report.docx'
    
    if os.path.exists(map_report_path):
        return send_file(map_report_path, as_attachment=True)
    else:
        return "Report not found", 404

@app.route('/download-city-report', methods=['GET'])
@login_required
def download_city_report():
    city_report_path = 'static/Top_10_Cities_Report.docx'
    
    if os.path.exists(city_report_path):
        return send_file(city_report_path, as_attachment=True)
    else:
        return "Report not found", 404

@app.route('/accident_investigation')
@login_required
def ACCIDENT_INVESTIGATION():
    return render_template('accident_investigation.html')

@app.route('/safety_audit_assessment')
@login_required
def SAFETYAUDITS_ASSESSMENTS():
    return render_template('safety_audit_assessment.html')

@app.route('/research_development')
@login_required
def RESEARCH_DEVELOPMENT():
    return render_template('research_development.html')

@app.route('/custom_solution_software')
@login_required
def CUSTOM_SOLUTION_SOFTWARE():
    return render_template('custom_solution_software.html')

@app.route('/safety_equipment')
@login_required
def SAFETY_EQUIPMENT():
    return render_template('safety_equipment.html')


@app.route('/events_calendar')
@login_required
def EVENTS():
    return render_template('events_calendar.html')

@app.route('/community_guidlines')
@login_required
def COMMUNITY_GUIDLINES():
    return render_template('community_guidlines.html')

if __name__ == '__main__':
    app.run(debug=True, port=3672)
